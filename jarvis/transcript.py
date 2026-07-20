"""Conversation transcript: log every exchange, and export it to a Word doc."""

from __future__ import annotations

import json
import os
from datetime import datetime

_DATA = os.path.join(os.path.dirname(os.path.dirname(__file__)), "data")
_PATH = os.path.join(_DATA, "conversation.jsonl")


def clear() -> None:
    """Start a fresh recording — wipe the previous transcript."""
    try:
        os.remove(_PATH)
    except FileNotFoundError:
        pass


def log(speaker: str, text: str) -> None:
    text = (text or "").strip()
    if not text:
        return
    os.makedirs(_DATA, exist_ok=True)
    entry = {"ts": datetime.now().isoformat(timespec="seconds"), "speaker": speaker, "text": text}
    with open(_PATH, "a") as fh:
        fh.write(json.dumps(entry) + "\n")


def _load() -> list:
    try:
        with open(_PATH) as fh:
            return [json.loads(line) for line in fh if line.strip()]
    except FileNotFoundError:
        return []


def export_docx(dest: str | None = None) -> str:
    """Write the full conversation to a Word document on the Desktop."""
    from docx import Document
    from docx.shared import Pt, RGBColor

    entries = _load()
    doc = Document()
    doc.add_heading("Jarvis — Conversation Log", level=0)
    doc.add_paragraph(datetime.now().strftime("Exported %A, %B %d, %Y at %I:%M %p"))

    last_date = None
    for e in entries:
        try:
            dt = datetime.fromisoformat(e["ts"])
        except Exception:
            dt = None
        day = dt.strftime("%A, %B %d, %Y") if dt else ""
        if day and day != last_date:
            doc.add_heading(day, level=2)
            last_date = day
        p = doc.add_paragraph()
        who = p.add_run(f"{e['speaker']}: ")
        who.bold = True
        if e["speaker"].lower().startswith("j"):
            who.font.color.rgb = RGBColor(0x6A, 0x4C, 0xD0)  # purple for Jarvis
        p.add_run(e["text"])

    if dest is None:
        stamp = datetime.now().strftime("%Y-%m-%d")
        dest = os.path.join(
            os.path.expanduser("~/Desktop"), f"Jarvis Conversation {stamp}.docx"
        )
    doc.save(dest)
    return dest
